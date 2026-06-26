"""Fill the KROS Word template from paper_ko.md WITH formatting applied.

Applies the JKROS typographic spec programmatically: per-element fonts/sizes,
alignment, first-line indent, line spacing, heading levels (대/중/소제목),
italic reference venues, superscript [n] citations, table/caption fonts.

Not automatable here (do by hand in Word): 장평(98%)/자간(-5), Equation Editor
objects, author-bio photo. Algorithms are emitted as a monospace listing box.

    python3 docs/build_docx.py     # -> /home/user/Downloads/MARS_JKROS_draft.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn

SAMPLE = "/home/user/Downloads/KROS_Revised_Sample_수정.docx"
OUT = "/home/user/Downloads/MARS_JKROS_draft.docx"
KO = Path(__file__).parent / "paper_ko.md"
FIGDIR = Path(__file__).parent.parent / "eval" / "figs"

# fonts
F_TITLE = "HY견명조"; F_HEAD = "HY태고딕"; F_ABS = "HY헤드라인M"
F_BODY = "바탕"; TNR = "Times New Roman"; ARIAL = "Arial"

KO_TITLE = "물류 로봇 fleet을 위한 LLM 감독 에이전트의 결정론적 검증: 다중 모델 연구"
EN_TITLE = "Deterministic Validation of LLM Supervisory Agents for Warehouse Robot Fleets: A Multi-Model Study"
KO_AUTHOR = "이 명 일"
EN_AUTHOR = "Myong-Il Lee"
KEYWORDS = "LLM agents, Retrieval-augmented generation, Deterministic validation, Warehouse robot fleet, AI safety"
AFFIL = "†Corresponding author: Student, Dept. of Cyber Security, Korea Polytechnic University Gangseo Campus, Seoul, Korea (2220110150@office.kopo.ac.kr)"
FIGS = [
    ("fig_mm_rag.png", "Fig. 1. Diagnosis cause accuracy by model, RAG on vs. off (test, n=100)."),
    ("fig_mm_safety.png", "Fig. 2. Confident-wrong rate by model, RAG on vs. off."),
    ("fig_mm_intent.png", "Fig. 3. Intent defense-in-depth by model over 15 unsafe intents."),
]


def setfont(run, name, size, bold=False, italic=False, sup=False):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.superscript = sup
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = rPr.makeelement(qn("w:rFonts"), {}); rPr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), name)


def para(doc, align=AL.JUSTIFY, indent=None, before=0, after=0, ls=1.0):
    p = doc.add_paragraph(); pf = p.paragraph_format
    p.alignment = align
    if indent is not None: pf.first_line_indent = indent
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = ls
    return p


def strip_md(t):
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"\*(.+?)\*", r"\1", t)
    t = t.replace("`", "")
    return t


def body_run(p, text):
    """Body text: Korean->바탕, Latin->TNR (set both), [n] citations superscript."""
    for tok in re.split(r"(\[\d+\])", strip_md(text)):
        if not tok:
            continue
        r = p.add_run(tok)
        if re.fullmatch(r"\[\d+\]", tok):
            setfont(r, TNR, 9.5, sup=True)
        else:
            setfont(r, F_BODY, 9.5)
            r._element.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:ascii"), TNR)
            r._element.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:hAnsi"), TNR)


def ref_runs(p, text):
    """Reference line: italicize *venue* spans, TNR 9."""
    for i, span in enumerate(re.split(r"\*(.+?)\*", text)):
        if not span:
            continue
        r = p.add_run(span); setfont(r, TNR, 9, italic=(i % 2 == 1))


def get_section(s, start, end):
    return "\n".join(KO.read_text().splitlines()).split(start, 1)[1].split(end, 1)[0]


def main():
    doc = Document(SAMPLE)

    # --- title block: replace sample placeholders, then re-format ---
    def setp(par, text):
        for r in list(par.runs):
            r.text = ""
        (par.runs[0] if par.runs else par.add_run("")).text = text
        return par
    src = KO.read_text()
    abstract = src.split("## Abstract\n", 1)[1].split("\n**Keywords")[0].strip()
    abstract = " ".join(l.strip() for l in abstract.split("\n") if l.strip())

    for p in doc.paragraphs[:14]:
        t = p.text.strip()
        if "로봇학회논문지 논문 작성 방법" in t:
            setp(p, KO_TITLE); p.alignment = AL.CENTER
            for r in p.runs: setfont(r, F_TITLE, 19)
        elif t.startswith("Preparation of Papers"):
            setp(p, EN_TITLE); p.alignment = AL.CENTER
            for r in p.runs: setfont(r, F_TITLE, 16)
        elif "홍 길 동" in t:
            p.text = ""; p.alignment = AL.CENTER
            r = p.add_run(KO_AUTHOR); setfont(r, F_TITLE, 11)
            r2 = p.add_run("†"); setfont(r2, F_TITLE, 11, sup=True)
        elif t.startswith("Gildong Hong"):
            p.text = ""; p.alignment = AL.CENTER
            r = p.add_run(EN_AUTHOR); setfont(r, F_TITLE, 11)
            r2 = p.add_run("†"); setfont(r2, F_TITLE, 11, sup=True)
        elif t.startswith("Abstract"):
            p.text = ""; p.alignment = AL.JUSTIFY
            r = p.add_run("Abstract  "); setfont(r, F_ABS, 8.5)
            r2 = p.add_run(abstract); setfont(r2, TNR, 10)
        elif t.startswith("Keywords"):
            p.text = ""; p.alignment = AL.JUSTIFY
            r = p.add_run("Keywords:  "); setfont(r, F_ABS, 8.5)
            r2 = p.add_run(KEYWORDS); setfont(r2, TNR, 10)
        elif t.startswith("※"):
            setp(p, "")
        elif t.startswith("1. Principal"):
            setp(p, AFFIL); p.alignment = AL.JUSTIFY
            for r in p.runs: setfont(r, ARIAL, 7.5)
        elif t.startswith(("2. Manager", "†Associate")):
            setp(p, "")

    # --- drop sample body + table ---
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("1.") and p.style.name == "hstyle1":
            start = i; break
    if start is not None:
        for p in list(doc.paragraphs)[start:]:
            p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # --- body ---
    lines = src.splitlines()
    bi = next(i for i, l in enumerate(lines) if l.startswith("## 1. 서론"))
    ei = next(i for i, l in enumerate(lines) if l.startswith("## Figures"))
    blk = lines[bi:ei]
    i = 0
    while i < len(blk):
        s = blk[i].strip()
        if not s:
            i += 1; continue
        if s.startswith("## "):       # 대제목
            p = para(doc, AL.CENTER, before=12, after=6)
            setfont(p.add_run(strip_md(s[3:])), F_HEAD, 12); i += 1; continue
        if s.startswith("### "):      # 중제목 (N.N)
            p = para(doc, AL.LEFT, before=10, after=10)
            setfont(p.add_run(strip_md(s[4:])), F_BODY, 10.5, bold=True); i += 1; continue
        if s.startswith("[FIG:"):     # inline architecture/algorithm image
            fn = s[5:].rstrip("]").strip()
            fp = FIGDIR / fn
            if fp.exists():
                pic = para(doc, AL.CENTER); pic.add_run().add_picture(str(fp), width=Inches(3.4))
            i += 1; continue
        if s.startswith("```"):       # trust formula listing (monospace)
            i += 1; buf = []
            while i < len(blk) and not blk[i].strip().startswith("```"):
                buf.append(blk[i]); i += 1
            i += 1
            p = para(doc, AL.LEFT)
            setfont(p.add_run("\n".join(buf)), "Courier New", 8)
            continue
        if s.startswith(">"):         # algorithm caption lines
            p = para(doc, AL.LEFT); setfont(p.add_run(strip_md(s.lstrip("> "))), TNR, 9)
            i += 1; continue
        if s.startswith("|"):         # table
            rows = []
            while i < len(blk) and blk[i].strip().startswith("|"):
                row = blk[i].strip()
                if not re.match(r"^\|[\s:|-]+\|?$", row):
                    rows.append([strip_md(c.strip()) for c in row.strip("|").split("|")])
                i += 1
            if rows:
                tb = doc.add_table(rows=len(rows), cols=len(rows[0])); tb.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, val in enumerate(row):
                        if c < len(tb.rows[r].cells):
                            cell = tb.rows[r].cells[c]; cell.text = ""
                            setfont(cell.paragraphs[0].add_run(val), TNR, 9, bold=(r == 0))
            continue
        if s.startswith("Table ") or s.startswith("**Table"):  # table caption
            p = para(doc, AL.LEFT); setfont(p.add_run(strip_md(s)), TNR, 9.5); i += 1; continue
        # normal body paragraph — JOIN wrapped lines until blank/special line
        buf = []
        while i < len(blk):
            cur = blk[i].strip()
            if (not cur or cur.startswith(("## ", "### ", "```", ">", "|", "Table ", "**Table"))):
                break
            if buf and re.match(r"^(\d+\.|-)\s", cur):   # new list item -> new paragraph
                break
            buf.append(cur); i += 1
        p = para(doc, AL.JUSTIFY, indent=Cm(0.35)); body_run(p, " ".join(buf))

    # --- figures ---
    p = para(doc, AL.CENTER, before=12, after=6); setfont(p.add_run("Figures"), F_HEAD, 12)
    for fn, cap in FIGS:
        fp = FIGDIR / fn
        if fp.exists():
            pic = para(doc, AL.CENTER); pic.add_run().add_picture(str(fp), width=Inches(3.0))
            cp = para(doc, AL.JUSTIFY); setfont(cp.add_run(cap), TNR, 9.5)

    # --- references ---
    rs = next(i for i, l in enumerate(lines) if l.strip() == "## References")
    p = para(doc, AL.CENTER, before=12, after=6); setfont(p.add_run("References"), F_HEAD, 12)
    for l in lines[rs+1:]:
        s = l.strip()
        if s.startswith("["):
            rp = para(doc, AL.JUSTIFY); ref_runs(rp, s)

    doc.save(OUT)
    print("saved", OUT, "| tables:", len(doc.tables))


if __name__ == "__main__":
    main()
